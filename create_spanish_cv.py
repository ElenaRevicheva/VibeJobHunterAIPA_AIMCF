#!/usr/bin/env python3
"""
Create Elena Revicheva's Spanish CV as a Word document with clickable links
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a w:rPr element
    rPr = OxmlElement('w:rPr')

    # Add color to the text
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    # Add underline to the text
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)

    # Create a w:t element
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink

def create_cv_document():
    """Create the Spanish CV Word document"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Elena Revicheva', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Ingeniera y Fundadora centrada en IA | Creando IA Emocionalmente Inteligente')
    run.bold = True
    run.font.size = Pt(14)
    
    # Contact info
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run('📍 Panama City, Panama (Remote) | 🌎 EN/ES | ')
    contact.add_run('📧 ')
    add_hyperlink(contact, 'E-mail', 'mailto:aipa@aideazz.xyz')
    contact.add_run(' | 📱 ')
    add_hyperlink(contact, 'WhatsApp', 'https://wa.me/50766623757')
    contact.add_run(' | ')
    add_hyperlink(contact, 'Telegram', 'https://t.me/ElenaRevicheva')
    
    # Links row
    links = doc.add_paragraph()
    links.alignment = WD_ALIGN_PARAGRAPH.CENTER
    links.add_run('🔗 ')
    add_hyperlink(links, 'Portfolio', 'https://aideazz.xyz/card')
    links.add_run(' | ')
    add_hyperlink(links, 'GitHub', 'https://github.com/ElenaRevicheva')
    links.add_run(' | ')
    add_hyperlink(links, 'LinkedIn', 'https://linkedin.com/in/elenarevicheva')
    links.add_run(' | ')
    add_hyperlink(links, 'Website', 'https://aideazz.xyz')
    links.add_run(' | 🌐 ENS: aideazz.eth')
    
    doc.add_paragraph()  # Space
    
    # Summary section
    doc.add_heading('💡 Resumen', 1)
    
    summary_quote = doc.add_paragraph()
    run = summary_quote.add_run('"Transformando sueños de equipos completos en creaciones individuales — impulsadas por el vibe coding."')
    run.italic = True
    
    doc.add_paragraph(
        'Fundadora de AIdeazz.xyz, dedicada a crear Asistentes Personales de IA Emocionalmente Inteligentes (AIPAs): '
        'compañeros conscientes diseñados para la educación, la adaptación cultural y el crecimiento personal.'
    )
    
    doc.add_paragraph(
        'Ex CEO y CLO en el sector de Gobierno Electrónico (Rusia). Me reubiqué en Panamá en 2022 para reconstruirme '
        'desde cero y lanzar 6 productos de IA desarrollados en solitario por menos de $15K.'
    )
    
    doc.add_paragraph(
        'Actualmente busco integrarme en una startup de IA como Ingeniera de IA / Product Builder / Ingeniera Fundadora.'
    )
    
    # Key Achievements
    doc.add_heading('📊 Logros Clave', 1)
    
    achievements = [
        '6 productos de IA (2 agentes de IA activos) en 7 meses — desarrollo full-stack en solitario (Python, TypeScript, React).',
        'Reducción del 98% en costos frente al desarrollo en equipo ($900K → <$15K).',
        'Usuarios en 19 países hispanohablantes (mercado bilateral); arquitectura bilingüe (EN/ES).',
        'Integración de más de 8 servicios de IA (Claude, GPT, Whisper, TTS, OCR, ElizaOS, HeyGen).',
        'Suscripciones PayPal activas; pagos en criptomonedas en fase de prueba.',
        'Constructora 0→1: Visión → Diseño → Desarrollo → Implementación → Crecimiento.'
    ]
    
    for achievement in achievements:
        p = doc.add_paragraph(achievement, style='List Bullet')
    
    # Technical Stack
    doc.add_heading('⚙️ Stack Técnico', 1)
    
    tech_stack = {
        'IA/ML': 'GPT · Claude · Whisper · TTS · MCP · LangChain · ElizaOS',
        'Lenguajes': 'Python · TypeScript · JavaScript · SQL',
        'Frameworks': 'React · Flask · Node.js · Vite',
        'Infraestructura': 'PostgreSQL · Supabase · Docker · Railway',
        'Frontend': 'Tailwind CSS · shadcn/ui · Framer Motion · i18next',
        'APIs': 'WhatsApp · Telegram · PayPal · Twitter · CCXT',
        'Web3': 'Polygon · Thirdweb · MetaMask · IPFS · Diseño DAO'
    }
    
    for category, technologies in tech_stack.items():
        p = doc.add_paragraph()
        run = p.add_run(f'{category}: ')
        run.bold = True
        p.add_run(technologies)
    
    # Products
    doc.add_heading('🚀 Productos del Ecosistema AIdeazz — Desarrollados en Solitario', 1)
    
    products_intro = doc.add_paragraph()
    run = products_intro.add_run('🧠 Fundadora e Ingeniera Principal — AIdeazz.xyz | Panamá | 2025–Presente')
    run.bold = True
    
    doc.add_paragraph('Desarrollo de Asistentes Personales de IA Emocionalmente Inteligentes (AIPAs).')
    
    # Product 1: EspaLuz
    product1 = doc.add_paragraph()
    run = product1.add_run('✅ EspaLuz – Tutor de Español con IA (ACTIVO)')
    run.bold = True
    
    doc.add_paragraph(
        'Tutor bilingüe que conecta a expatriados y locales (EN↔ES). Memoria emocional persistente, OCR, TTS y síntesis de voz.'
    )
    
    p = doc.add_paragraph('→ Activo en WhatsApp y Telegram; usuarios en 19 países; suscripciones PayPal habilitadas.')
    p = doc.add_paragraph('🔗 ')
    add_hyperlink(p, 'EspaLuz WhatsApp', 'https://wa.me/50766623757')
    p.add_run(' | ')
    add_hyperlink(p, 'EspaLuz Telegram', 'https://t.me/EspaLuzFamily_bot')
    p.add_run(' | ')
    add_hyperlink(p, 'EspaLuz SaaS Web App', 'https://espaluz-ai-language-tutor.lovable.app/')
    
    # Product 2: ALGOM Alpha
    product2 = doc.add_paragraph()
    run = product2.add_run('✅ ALGOM Alpha – Mentor Cripto con IA (ACTIVO)')
    run.bold = True
    
    p = doc.add_paragraph('🔗 ')
    add_hyperlink(p, 'Algom Alpha on X', 'https://x.com/algom_alpha')
    
    doc.add_paragraph(
        'Enseña trading seguro y alfabetización digital mediante operaciones simuladas autónomas. '
        'Tecnología: Node.js, ElizaOS, CCXT, API de Twitter.'
    )
    
    # Product 3: Influencer Bot
    product3 = doc.add_paragraph()
    run = product3.add_run('✅ EspaLuz Influencer Bot (ACTIVO)')
    run.bold = True
    
    p = doc.add_paragraph('🔗 ')
    add_hyperlink(p, 'EspaLuz Influencer', 'https://t.me/Influencer_EspaLuz_bot')
    
    doc.add_paragraph('Automatiza la generación de contenido con IA y publicaciones en LinkedIn + Instagram vía Buffer.')
    
    # Product 4: Atuona
    product4 = doc.add_paragraph()
    run = product4.add_run('✅ Atuona NFT Gallery (ACTIVO)')
    run.bold = True
    
    p = doc.add_paragraph('🔗 ')
    add_hyperlink(p, 'atuona.xyz', 'https://atuona.xyz')
    
    doc.add_paragraph(
        'Lanzamientos de poesía NFT con enfoque en mindfulness en Polygon. '
        'Stack: Thirdweb SDK, React, IPFS.'
    )
    
    # Product 5: Main Website
    product5 = doc.add_paragraph()
    run = product5.add_run('✅ Sitio Principal de AIdeazz — Plataforma del Ecosistema (ACTIVO)')
    run.bold = True
    
    p = doc.add_paragraph('🔗 ')
    add_hyperlink(p, 'aideazz.xyz', 'https://aideazz.xyz')
    
    doc.add_paragraph(
        'UX bilingüe (EN/ES), diseño responsivo, construido con React, Tailwind y Framer Motion.'
    )
    
    # Previous Experience
    doc.add_heading('🧩 Experiencia Previa', 1)
    
    exp1 = doc.add_paragraph()
    run = exp1.add_run('Cofundadora Operativa — OmniBazaar Marketplace Startup | Remoto | 2024–2025')
    run.bold = True
    doc.add_paragraph('Estructuración de DAO LLC (Islas Marshall), tokenomics y modelo de gobernanza.')
    
    exp2 = doc.add_paragraph()
    run = exp2.add_run('Subdirectora General y CLO — JSC "E-GOV OPERATOR" | Rusia | 2011–2018')
    run.bold = True
    doc.add_paragraph('Lideré la transformación digital regional de los servicios públicos. Gestión de TI, RRHH y cumplimiento normativo.')
    
    exp3 = doc.add_paragraph()
    run = exp3.add_run('Subdirectora General (Desarrollo de Negocios) — Fundery LLC | Rusia | 2017–2018')
    run.bold = True
    doc.add_paragraph('Cumplimiento de ICO y relaciones con inversionistas durante el auge blockchain.')
    
    # Education
    doc.add_heading('🎓 Educación y Certificaciones', 1)
    
    education = [
        'Polkadot Blockchain Academy, PBA-X Wave #3 (curso en línea, 2025)',
        'How-To-DAO Cohort Graduate (curso en línea, 2025)',
        'M.A. en Psicología Social, Universidad Estatal de Penza (Rusia, 2018)',
        'Regulación Blockchain, MGIMO (Moscú, 2017)',
        'Programa Presidencial de Gestión Ejecutiva, RANEPA (Moscú, 2015)',
        '— Pasantía en Nyskapingsparken Innovation Park, Bergen, Noruega'
    ]
    
    for edu in education:
        p = doc.add_paragraph(edu, style='List Bullet')
    
    # Languages
    doc.add_heading('🌍 Idiomas', 1)
    doc.add_paragraph('🇷🇺 Ruso (Nativo) | 🇬🇧 Inglés (Fluido) | 🇪🇸 Español (Intermedio) | 🇫🇷 Francés (Básico)')
    
    # Open To
    doc.add_heading('💼 Abierta a Roles de Tiempo Completo o Parcial', 1)
    
    roles = doc.add_paragraph()
    run = roles.add_run('✅ AI Product Manager | Full-Stack AI Engineer | Founding Engineer')
    run.bold = True
    
    roles2 = doc.add_paragraph()
    run = roles2.add_run('✅ LLM Engineer | AI Solutions Architect | AI Growth Engineer')
    run.bold = True
    
    roles3 = doc.add_paragraph()
    run = roles3.add_run('✅ Enfoque híbrido: Rol + Inversión Pre-seed para AIdeazz (ejecución paralela).')
    run.bold = True
    
    # Why Work With Me
    doc.add_heading('🌟 Por qué Colaborar Conmigo', 1)
    
    doc.add_paragraph(
        'Ejecución de nivel fundadora unida a una visión de IA emocional — del concepto al GTM. '
        'Nativa Web3 y bilingüe, creo la próxima generación de IA que crece con los humanos y evoluciona a lo largo de su camino.'
    )
    
    # Save document
    output_path = '/workspace/Elena_Revicheva_CV_Spanish.docx'
    doc.save(output_path)
    print(f"✅ Spanish CV Word document created: {output_path}")
    return output_path

if __name__ == '__main__':
    create_cv_document()
